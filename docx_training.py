import docx     # Модуль работы с документами Word
import re

my_text = docx.Document('C:\Training\ПЗ 0503760 071.docx')              # Загружаем документ
my_text_1 = docx.Document('C:\Training\kest.docx')
#print(len(my_text.tables))                                              # Выводим размеры табЛицы в документе Word

paragraph_size = len(my_text.paragraphs)
paragraph_size_1 = len(my_text_1.paragraphs)
print(paragraph_size, paragraph_size_1)


paragraphs = []
for paragraph in my_text.paragraphs:
    paragraphs.append(paragraph.text)
pattern = re.compile(r'Раздел.*')
#pattern = re.match(r'Раздел.*', paragraphs)
match = list(filter(pattern.match, paragraphs))
    #print('\n'.join(paragraphs))
print('\n'.join(match))

'''
# Работа непосредственно с текстом 
text = []
for paragraph in my_text.paragraphs:
    text.append(paragraph.text)
print('\n'.join(text))
'''


'''
table_buh_ot = my_text.tables[12]
table_size = len(my_text.tables[12].rows)
n_form = table_buh_ot.rows[-1].cells[0].text
name_form = table_buh_ot.rows[-1].cells[1].text
code_form = table_buh_ot.rows[-1].cells[2].text
analitic_form = table_buh_ot.rows[-1].cells[3].text
analitic_size_form = table_buh_ot.rows[-1].cells[4].text

print(n_form, name_form, code_form, analitic_form, analitic_size_form, sep='; ')
print(table_size)

for row in table_buh_ot.rows:
    #dataframe = ''
    dataframe = []
    for cell in row.cells:
        #dataframe += cell.text + ','
        dataframe.append(cell.text)
        dataframe = list(map(lambda s: s.replace(u'\xa0', u'').strip('\n'), dataframe))
        #dataframe.replace(u'\xa0', u' ')
        #print(dataframe)
'''

#my_text.save('C:\Training\kest.docx')




